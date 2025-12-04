from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from backend.utils.compute_file_hash import compute_file_hash
from documents.models import Document
from django.core.files.storage import default_storage
from backend.utils.qdrant_client import (
    get_qdrant_client, get_sparse_builder, generate_sparse_vector,
    SparseVectorBuilder
)
from qdrant_client.http.models import (
    PointStruct, VectorParams, Distance, Filter, FieldCondition, MatchValue,
    SparseVectorParams, SparseIndexParams, NamedSparseVector
)
from transformers import AutoTokenizer, AutoModel
from backend.utils.embeddings import HuggingfaceEmbeddingsModel
import os
import uuid
import re
import nltk
import fitz
from typing import List, Dict, Any, Optional, Tuple
from django.conf import settings
import logging
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from backend.utils.qdrant_client import QdrantClient

# Chunking and evaluation
from ..tools.chunk_eval import eval_chunks
from ..tools.export_chunks import write_chunks_csv
from ..tools.chunking import build_page_hierarchy_map

# MongoDB imports
from backend.utils.mongo_repository import MongoRepository
from backend.utils.mongo_models import MongoDocument, MongoChunk

logger = logging.getLogger(__name__)


# ============================================================================
# DOCUMENT HIERARCHY EXTRACTION
# ============================================================================

def extract_pdf_hierarchy(pdf_path: str) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """
    Extract document structure (chapters, sections) from PDF using TOC.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Tuple of:
        - List of section dicts: [{title, level, page}, ...]
        - Document title (from TOC or first heading)
    """
    sections = []
    document_title = None
    
    try:
        with fitz.open(pdf_path) as doc:
            # Try to get TOC (table of contents)
            toc = doc.get_toc()  # Returns [[level, title, page], ...]
            
            if toc:
                for item in toc:
                    level = item[0]  # Heading level (1, 2, 3...)
                    title = item[1]  # Heading title
                    page = item[2] - 1  # Convert to 0-indexed
                    
                    sections.append({
                        "title": title.strip(),
                        "level": level,
                        "page": max(0, page),  # Ensure non-negative
                    })
                    
                    # First level-1 heading is likely the document title
                    if level == 1 and document_title is None:
                        document_title = title.strip()
                
                logger.info(f"Extracted {len(sections)} sections from PDF TOC")
            else:
                # No TOC - try to detect structure from text
                sections = _detect_structure_from_pdf_text(doc)
                
                if sections:
                    # Get document title from first major heading
                    for s in sections:
                        if s.get("level", 99) <= 1:
                            document_title = s.get("title")
                            break
                    logger.info(f"Detected {len(sections)} sections from PDF text (no TOC)")
                else:
                    logger.info("No document structure detected in PDF")
                    
    except Exception as e:
        logger.warning(f"Error extracting PDF hierarchy: {e}")
    
    return sections, document_title


def _detect_structure_from_pdf_text(doc) -> List[Dict[str, Any]]:
    """
    Detect document structure from PDF text when no TOC is available.
    Uses regex patterns to identify chapter/section headings.
    """
    sections = []
    
    # Patterns for Vietnamese and English headings
    patterns = [
        # Vietnamese chapter patterns
        (r'^(?:Chương|CHƯƠNG)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Phần|PHẦN)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        # Vietnamese section patterns
        (r'^(?:Mục|MỤC)\s+(\d+(?:\.\d+)*)[\.:]\s*(.+)$', 2),
        (r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)$', 2),  # Numbered sections like "1.1 Title"
        # English patterns
        (r'^(?:Chapter|CHAPTER)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Part|PART)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Section|SECTION)\s+(\d+(?:\.\d+)*)[\.:]\s*(.+)$', 2),
        # All-caps headings (likely section titles)
        (r'^([A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤƯỨỪỬỮỰ\s]{15,})$', 1),
    ]
    
    for page_num, page in enumerate(doc):
        text = page.get_text()
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            for pattern, level in patterns:
                match = re.match(pattern, line, re.MULTILINE)
                if match:
                    # Get the title (last group or full match for all-caps)
                    title = match.group(2) if match.lastindex >= 2 else line
                    title = title.strip()
                    
                    if title and len(title) > 2:
                        sections.append({
                            "title": title,
                            "level": level,
                            "page": page_num,
                        })
                    break
    
    # Deduplicate sections on same page with same level
    seen = set()
    unique_sections = []
    for s in sections:
        key = (s["page"], s["level"], s["title"][:30])
        if key not in seen:
            seen.add(key)
            unique_sections.append(s)
    
    return unique_sections


def extract_docx_hierarchy(docx_path: str) -> Tuple[List[Dict[str, Any]], Optional[str]]:
    """
    Extract document structure from DOCX using paragraph styles.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Tuple of:
        - List of section dicts: [{title, level, paragraph_index}, ...]
        - Document title (from Title style or first heading)
    """
    from docx import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE
    
    sections = []
    document_title = None
    
    try:
        doc = DocxDocument(docx_path)
        
        # Map style names to hierarchy levels
        style_levels = {
            "Title": 0,
            "Heading 1": 1,
            "Heading 2": 2,
            "Heading 3": 3,
            "Heading 4": 4,
            "Heading 5": 5,
            "Heading 6": 6,
            # Vietnamese style names
            "Tiêu đề": 0,
            "Tiêu đề 1": 1,
            "Tiêu đề 2": 2,
            "Tiêu đề 3": 3,
        }
        
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else ""
            
            if style_name in style_levels:
                level = style_levels[style_name]
                
                sections.append({
                    "title": text,
                    "level": level if level > 0 else 1,  # Treat Title as level 1 for hierarchy
                    "page": idx // 30,  # Approximate page (30 paragraphs per page)
                })
                
                # Get document title from Title style or first Heading 1
                if level == 0 and document_title is None:
                    document_title = text
                elif level == 1 and document_title is None:
                    document_title = text
        
        if sections:
            logger.info(f"Extracted {len(sections)} sections from DOCX styles")
        else:
            # Fallback: detect structure from text patterns
            sections = _detect_structure_from_docx_text(doc)
            logger.info(f"Detected {len(sections)} sections from DOCX text (no styles)")
            
    except Exception as e:
        logger.warning(f"Error extracting DOCX hierarchy: {e}")
    
    return sections, document_title


def _detect_structure_from_docx_text(doc) -> List[Dict[str, Any]]:
    """
    Detect document structure from DOCX text when styles aren't used.
    """
    sections = []
    
    # Same patterns as PDF
    patterns = [
        (r'^(?:Chương|CHƯƠNG)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Phần|PHẦN)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Mục|MỤC)\s+(\d+(?:\.\d+)*)[\.:]\s*(.+)$', 2),
        (r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)$', 2),
        (r'^(?:Chapter|CHAPTER)\s+(\d+|[IVXLCDM]+)[\.:]\s*(.+)$', 1),
        (r'^(?:Section|SECTION)\s+(\d+(?:\.\d+)*)[\.:]\s*(.+)$', 2),
    ]
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 3:
            continue
        
        for pattern, level in patterns:
            match = re.match(pattern, text, re.MULTILINE)
            if match:
                title = match.group(2) if match.lastindex >= 2 else text
                title = title.strip()
                
                if title and len(title) > 2:
                    sections.append({
                        "title": title,
                        "level": level,
                        "page": idx // 30,
                    })
                break
    
    return sections


# ============================================================================
# UPLOAD VIEW FUNCTIONS
# ============================================================================

@csrf_exempt
@require_http_methods(["POST"])
def upload_document(request):
    """
    Handle document upload and processing.
    """
    try:
        force = request.GET.get('force', 'false').lower() == 'true'
        
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        qdrant_client = QdrantClient()
        
        # Now initialize/create collection
        qdrant_client.init_collection()
        
        # Process the file
        chunks = chunk_document(file)
        
        if not chunks:
            return JsonResponse({'error': 'No content extracted from file'}, status=400)
        
        # Generate embeddings and store in Qdrant
        texts = [chunk['text'] for chunk in chunks]
        embeddings = qdrant_client.embed_texts(texts)
        
        # Upload to vector DB
        qdrant_client.upload_vectors(embeddings, chunks)
        
        logger.info(f"Successfully uploaded document: {file.name} with {len(chunks)} chunks")
        
        return JsonResponse({
            'message': 'Document uploaded successfully',
            'file_name': file.name,
            'chunks_count': len(chunks)
        })
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)

class FileUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    
    # Get embedding model from env or use default (dangvantuan/vietnamese-embedding, 768-dim)
    model_name = os.getenv("EMBEDDING_MODEL", "dangvantuan/vietnamese-embedding")
    
    # Initialize tokenizer - handle both formats
    try:
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        logger.info(f"Loaded tokenizer for model: {model_name}")
    except Exception as e:
        logger.warning(f"Could not load tokenizer for {model_name}: {e}. Using default.")
        tokenizer = AutoTokenizer.from_pretrained("dangvantuan/vietnamese-embedding")
    
    # Initialize embeddings model using Transformers (HuggingfaceEmbeddingsModel)
    try:
        embeddings_model = HuggingfaceEmbeddingsModel(model_name)
        logger.info(f"Loaded HuggingfaceEmbeddingsModel with model: {model_name}")
    except Exception as e:
        logger.error(f"Failed to load HuggingfaceEmbeddingsModel {model_name}: {e}")
        embeddings_model = HuggingfaceEmbeddingsModel("dangvantuan/vietnamese-embedding")
        logger.info("Loaded default HuggingfaceEmbeddingsModel")
    
    mongo_repo = MongoRepository()  # MongoDB repository
    
    def delete_points_for_document(self, client, collection, document_id):
        """Delete all vector points for a specific document_id"""
        flt = Filter(must=[FieldCondition(key="document_id", match=MatchValue(value=document_id))])
        client.delete(collection_name=collection, points_selector=flt)
        logger.info(f"Deleted existing vector points for document_id: {document_id}")

    def post(self, request, *args, **kwargs):
        # Ensure NLTK 'punkt' is available
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')

        file = request.FILES.get('file')
        if not file:
            return Response(
                {"error": "No file uploaded. Make sure you use Body -> form-data and key name 'file' (type: File)."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Log file details
        logger.debug(f"Received file: name={file.name}, content_type={getattr(file, 'content_type', None)}, size={file.size}")

        # Basic validation
        allowed = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
        ]
        if file.content_type not in allowed:
            return Response(
                {"error": "Unsupported file type", "content_type": file.content_type},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Save the file temporarily
        saved_path = default_storage.save(f"uploads/{file.name}", file)
        logger.debug(f"Saved uploaded file to: {saved_path}")

        # Compute hash from the saved file
        saved_file_path = default_storage.path(saved_path) if hasattr(default_storage, 'path') else os.path.join(settings.MEDIA_ROOT, saved_path)
        with open(saved_file_path, "rb") as f:
            file_hash = compute_file_hash(f)

        # Check for force parameter - allows reingestion of the same file
        force = request.query_params.get("force") == "true"
        allow_duplicate = request.query_params.get("allow_duplicate", "false").lower() == "true"

        # Check for duplicates in Django DB
        existing_doc = Document.objects.filter(hash=file_hash).first()
        existing_mongo_doc = self.mongo_repo.get_document_by_hash(file_hash) if not force else None

        # Default: if document exists and no force/allow_duplicate, reject
        if existing_doc and not force and not allow_duplicate:
            default_storage.delete(saved_path)  # Clean up the temporary file
            return Response({
                'document_id': existing_doc.document_id,
                'filename': existing_doc.filename,
                'mimetype': existing_doc.mimetype,
                'size': existing_doc.size,
                'hash': existing_doc.hash,
                'message': 'Document already exists'
            })

        # If forcing, reuse the same document_id (hash) and delete existing points only
        if existing_doc and force:
            client = get_qdrant_client()
            collection_name = settings.QDRANT_COLLECTION
            # Remove ONLY this document's vector points from Qdrant (preserve other documents)
            try:
                self.delete_points_for_document(client, collection_name, existing_doc.document_id)
                logger.info(f"Deleted Qdrant points for document_id: {existing_doc.document_id}")
            except Exception as e:
                logger.warning(f"Could not delete existing Qdrant points: {str(e)}")

            # Delete existing chunks in MongoDB
            self.mongo_repo.delete_chunks_by_document(existing_doc.document_id)
            
            # Reset counters on the existing document row
            existing_doc.token_count = 0
            existing_doc.vector_dim = None
            existing_doc.save()
            document = existing_doc
            logger.info(f"Force reingestion of document: {document.document_id}")
        elif existing_doc and allow_duplicate:
            # Create a new document entry even though content hash exists (duplicate allowed)
            new_doc_id = uuid.uuid4().hex
            document = Document.objects.create(
                document_id=new_doc_id,
                filename=file.name,
                mimetype=file.content_type,
                size=file.size,
                hash=file_hash,
            )
            logger.info(f"Created new document entry for duplicate content: {document.document_id}")
        else:
            # Save metadata (new document)
            document = Document.objects.create(
                document_id=file_hash,
                filename=file.name,
                mimetype=file.content_type,
                size=file.size,
                hash=file_hash,
            )
        
        # Create MongoDB document entry (will be updated after processing)
        mongo_doc = MongoDocument(
            document_id=document.document_id,
            filename=file.name,
            content_hash=file_hash,
            mimetype=file.content_type,
            size_bytes=file.size,
            status="processing"
        )
        self.mongo_repo.save_document(mongo_doc)

        try:
            # Extract text and metadata based on file type
            if file.content_type == 'application/pdf':
                raw_blocks = self.extract_text_and_metadata_from_pdf(saved_file_path, document.document_id)
            elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                raw_blocks = self.extract_text_and_metadata_from_docx(saved_file_path, document.document_id)
            elif file.content_type == 'text/plain':
                raw_blocks = self.extract_text_and_metadata_from_txt(saved_file_path, document.document_id)
            else:
                raise ValueError("Unsupported file type")

            # Clean and chunk text with improved semantic chunking (P0 fix)
            from ..tools.chunking import clean_and_chunk_text, semantic_chunk_text
            from api.utils.query_preprocessing import preprocess_document_text
            
            # Apply consistent text preprocessing
            for block in raw_blocks:
                if 'text' in block and block['text']:
                    block['text'] = preprocess_document_text(block['text'])
            
            # Get structure map and document title from extraction
            structure_map = getattr(self, '_last_structure_map', None)
            document_title = getattr(self, '_last_document_title', None)
            
            # Use semantic chunking if enabled, otherwise use standard chunking
            use_semantic = os.getenv("USE_SEMANTIC_CHUNKING", "true").lower() == "true"
            
            if use_semantic:
                logger.info("Using semantic chunking with sentence boundaries and hierarchy")
                chunks = semantic_chunk_text(
                    raw_blocks, 
                    self.tokenizer,
                    target_tokens=int(os.getenv("CHUNK_TARGET_TOKENS", "400")),
                    max_tokens=int(os.getenv("CHUNK_MAX_TOKENS", "512")),
                    overlap_tokens=int(os.getenv("CHUNK_OVERLAP_TOKENS", "50")),
                    structure_map=structure_map,
                    document_title=document_title
                )
            else:
                logger.info("Using standard chunking with hierarchy")
                chunks = clean_and_chunk_text(
                    raw_blocks, 
                    self.tokenizer,
                    structure_map=structure_map,
                    document_title=document_title
                )
            
            logger.debug(f"Chunks generated: {chunks}")
            logger.debug(f"Number of chunks: {len(chunks)}")
            
            # Filter out chunks with empty or whitespace-only text
            original_count = len(chunks)
            chunks = [chunk for chunk in chunks if chunk.get('text') and chunk['text'].strip()]
            if len(chunks) != original_count:
                logger.warning(f"Filtered out {original_count - len(chunks)} empty chunks")
            
            if not chunks:
                raise ValueError("No valid text content extracted from the document")

            # Evaluate chunking quality before upload
            try:
                chunk_stats = eval_chunks(chunks, self.tokenizer, max_tokens=512)
                logger.info(f"Chunk evaluation: {chunk_stats}")
            except Exception as eval_exc:
                logger.exception("Chunk evaluation failed")
                chunk_stats = {"error": str(eval_exc)}

            # Generate embeddings using Transformers model
            chunk_texts = [chunk['text'] for chunk in chunks]
            
            # Debug: Log chunk texts before embedding
            logger.info(f"Processing {len(chunk_texts)} chunks for embedding")
            for i, text in enumerate(chunk_texts[:5]):  # Log first 5
                logger.debug(f"Chunk {i}: len={len(text)}, preview='{text[:100]}...'")
            
            # Validate and filter chunk texts
            valid_chunks = []
            valid_texts = []
            for i, (chunk, text) in enumerate(zip(chunks, chunk_texts)):
                if text and isinstance(text, str) and text.strip():
                    cleaned = text.strip()
                    if len(cleaned) > 0:
                        valid_chunks.append(chunk)
                        valid_texts.append(cleaned)
                    else:
                        logger.warning(f"Skipping empty chunk {i}")
                else:
                    logger.warning(f"Skipping invalid chunk {i}: type={type(text)}, value={repr(text)[:50]}")
            
            chunks = valid_chunks
            chunk_texts = valid_texts
            
            if not chunk_texts:
                raise ValueError("No valid text chunks to embed")
            
            logger.info(f"Embedding {len(chunk_texts)} valid chunks")
            
            try:
                embeddings = self.embeddings_model.encode(chunk_texts, convert_to_numpy=True)
            except Exception as e:
                logger.error(f"Embedding failed: {e}")
                # Debug: Try embedding one by one
                logger.info("Attempting individual chunk embedding to find problem...")
                for i, text in enumerate(chunk_texts):
                    try:
                        self.embeddings_model.encode([text], convert_to_numpy=True)
                        logger.debug(f"Chunk {i} OK")
                    except Exception as chunk_err:
                        logger.error(f"Chunk {i} FAILED: {chunk_err}")
                        logger.error(f"Problem text: {repr(text)}")
                raise

            # Ensure Qdrant collection exists
            client = get_qdrant_client()
            collection_name = settings.QDRANT_COLLECTION
            
            # Check if collection exists and determine if it's hybrid
            is_hybrid = False
            try:
                collection_info = client.get_collection(collection_name)
                # Check if sparse vectors are configured
                is_hybrid = (
                    hasattr(collection_info.config, 'sparse_vectors_config') and 
                    collection_info.config.sparse_vectors_config is not None
                )
                # Check vector config type
                vectors_config = collection_info.config.params.vectors
                if isinstance(vectors_config, dict):
                    has_dense_named = "dense" in vectors_config
                else:
                    has_dense_named = False
                    
                logger.info(f"Collection {collection_name} exists (hybrid={is_hybrid}, dense_named={has_dense_named})")
            except Exception:
                # Collection doesn't exist - create it
                try:
                    # Check if hybrid search is enabled
                    enable_hybrid = os.getenv("ENABLE_HYBRID_SEARCH", "true").lower() == "true"
                    vector_size = embeddings[0].shape[0] if len(embeddings) > 0 else int(os.getenv("EMBEDDING_DIM", "768"))
                    
                    if enable_hybrid:
                        # Create hybrid collection with dense + sparse vectors
                        client.create_collection(
                            collection_name=collection_name,
                            vectors_config={
                                "dense": VectorParams(
                                    size=vector_size,
                                    distance=Distance.COSINE
                                )
                            },
                            sparse_vectors_config={
                                "sparse": SparseVectorParams(
                                    index=SparseIndexParams(on_disk=False)
                                )
                            }
                        )
                        is_hybrid = True
                        has_dense_named = True
                        logger.info(f"Created HYBRID collection {collection_name} with dense ({vector_size}-dim) + sparse vectors")
                    else:
                        # Create legacy collection with named "default" vector
                        client.create_collection(
                            collection_name=collection_name,
                            vectors_config={
                                "default": VectorParams(
                                    size=vector_size,
                                    distance=Distance.COSINE
                                )
                            }
                        )
                        has_dense_named = False
                        logger.info(f"Created collection {collection_name} with vector size: {vector_size}")
                except Exception as e:
                    msg = str(e).lower()
                    if "already exists" in msg or "409" in msg:
                        logger.info(f"Collection {collection_name} already exists (ignored)")
                        is_hybrid = os.getenv("ENABLE_HYBRID_SEARCH", "true").lower() == "true"
                        has_dense_named = is_hybrid
                    else:
                        logger.exception(f"Failed to create collection {collection_name}")
                        raise

            # Generate sparse vectors if hybrid collection
            sparse_vectors = []
            if is_hybrid:
                try:
                    sparse_builder = get_sparse_builder()
                    
                    # Check if vocabulary is fitted
                    if not sparse_builder.vocabulary:
                        # Fit on current document chunks (bootstrap)
                        logger.info("Fitting sparse vector builder on current document...")
                        sparse_builder.fit(chunk_texts)
                        # Save for future use
                        cache_path = os.getenv("SPARSE_VOCAB_CACHE", "media/bm25_cache.pkl")
                        sparse_builder.save_vocabulary(cache_path)
                    
                    # Generate sparse vectors for all chunks
                    sparse_vectors = sparse_builder.transform_batch(chunk_texts)
                    logger.info(f"Generated {len([v for v in sparse_vectors if v is not None])} sparse vectors")
                except Exception as e:
                    logger.warning(f"Failed to generate sparse vectors: {e}. Continuing with dense-only.")
                    sparse_vectors = [None] * len(chunks)
            else:
                sparse_vectors = [None] * len(chunks)

            # Upload chunks and embeddings to Qdrant
            # Determine vector field name
            dense_vector_name = "dense" if (is_hybrid or has_dense_named) else "default"
            
            points = []
            for i, (embedding, chunk) in enumerate(zip(embeddings, chunks)):
                point_id = str(uuid.uuid4())
                
                # Build vectors dict
                vectors = {dense_vector_name: embedding.tolist()}
                
                # Add sparse vector if available
                if is_hybrid and i < len(sparse_vectors) and sparse_vectors[i] is not None:
                    sparse_vec = sparse_vectors[i]
                    vectors["sparse"] = {
                        "indices": sparse_vec.indices,
                        "values": sparse_vec.values
                    }
                
                # Clean up payload - ensure JSON serializable
                payload = {
                    "chunk_id": chunk.get("chunk_id"),
                    "document_id": chunk.get("document_id"),
                    "text": chunk.get("text"),
                    "page": chunk.get("page"),
                    "end_page": chunk.get("end_page"),
                    "section": chunk.get("section"),
                    "section_title": chunk.get("section_title"),
                    "section_path": chunk.get("section_path", []),
                    "hierarchy_level": chunk.get("hierarchy_level", 0),
                    "parent_section": chunk.get("parent_section"),
                    "document_title": chunk.get("document_title"),
                    "order_index": chunk.get("order_index"),
                    "text_len": chunk.get("text_len"),
                }
                
                points.append(PointStruct(
                    id=point_id,
                    vector=vectors,
                    payload=payload
                ))
            
            logger.debug(f"Uploading {len(points)} points to Qdrant (hybrid={is_hybrid})")

            response = client.upsert(
                collection_name=collection_name,
                points=points
            )
            logger.debug(f"Upsert response: {response}")

            # Save chunks to MongoDB in smaller batches
            mongo_chunks = []
            for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                mongo_chunk = MongoChunk(
                    chunk_id=chunk.get('chunk_id', f"{document.document_id}:{idx}"),
                    document_id=document.document_id,
                    text=chunk['text'],
                    embedding=embedding.tolist(),  # Convert numpy array to list
                    page=chunk.get('page'),
                    end_page=chunk.get('end_page'),
                    section=chunk.get('section'),
                    order_index=chunk.get('order_index', idx),
                    text_len=len(chunk['text']),
                    metadata={
                        "section_title": chunk.get("section_title"),
                        "section_path": chunk.get("section_path", []),
                        "hierarchy_level": chunk.get("hierarchy_level", 0),
                        "parent_section": chunk.get("parent_section"),
                        "document_title": chunk.get("document_title"),
                    },
                )
                mongo_chunks.append(mongo_chunk)
            
            # Bulk save chunks to MongoDB in batches of 100 to avoid BSON size limits
            batch_size = 100
            chunks_saved = 0
            for i in range(0, len(mongo_chunks), batch_size):
                batch = mongo_chunks[i:i + batch_size]
                try:
                    # Primary attempt: use repository helper
                    batch_count = self.mongo_repo.save_chunks_batch(batch)
                    # If repo returns falsy/0, fall back to direct insert into collection
                    if not batch_count:
                        logger.warning(f"repo.save_chunks_batch returned {batch_count} for batch {i // batch_size + 1}, falling back to direct insert")
                        docs = []
                        for c in batch:
                            docs.append({
                                "chunk_id": getattr(c, "chunk_id", None),
                                "document_id": getattr(c, "document_id", None),
                                "text": getattr(c, "text", None),
                                "embedding": getattr(c, "embedding", None),
                                "page": getattr(c, "page", None),
                                "end_page": getattr(c, "end_page", None),
                                "section": getattr(c, "section", None),
                                "order_index": getattr(c, "order_index", None),
                                "text_len": getattr(c, "text_len", None),
                                "metadata": getattr(c, "metadata", {}),
                            })
                        try:
                            result = self.mongo_repo.db["chunks"].insert_many(docs)
                            batch_count = len(result.inserted_ids)
                            logger.info(f"Fallback: inserted {batch_count} chunks directly into 'chunks' collection")
                        except Exception as ie:
                            logger.error(f"Fallback insert_many failed for batch {i // batch_size + 1}: {ie}")
                            batch_count = 0
                    chunks_saved += batch_count
                    logger.info(f"Saved batch {i // batch_size + 1}: {batch_count} / {len(batch)} chunks to MongoDB (total: {chunks_saved})")
                except Exception as e:
                    logger.exception(f"Failed to save batch {i // batch_size + 1}: {str(e)}")

            logger.info(f"✓ Total saved: {chunks_saved} chunks to MongoDB")

            # Update document status and store counts (Django DB)
            document.token_count = sum(len(self.tokenizer.encode(chunk['text'], add_special_tokens=False)) for chunk in chunks)
            document.vector_dim = embeddings[0].shape[0] if len(embeddings) > 0 else int(os.getenv("EMBEDDING_DIM", "768"))
            document.save()
            
            # Update MongoDB document with final status
            self.mongo_repo.update_document_status(
                document.document_id,
                status="processed",
                token_count=document.token_count,
                vector_dim=document.vector_dim,
                num_chunks=len(chunks)
            )

            # Write per-document CSV of chunks to MEDIA_ROOT/exports/{document_id}_chunks.csv
            try:
                csv_rel = f"exports/{document.document_id}_chunks.csv"
                csv_abs = os.path.join(settings.MEDIA_ROOT, csv_rel)
                # ensure exports directory exists
                os.makedirs(os.path.dirname(csv_abs), exist_ok=True)
                write_chunks_csv(chunks, csv_abs)
                logger.info(f"Wrote per-document chunks CSV: {csv_rel}")
            except Exception as csv_exc:
                logger.exception(f"Failed to write per-document CSV for {document.document_id}: {str(csv_exc)}")

            # Clean up the temporary file
            default_storage.delete(saved_path)

            # return chunk evaluation stats for quick validation (remove in production)
            return Response({"message": "File uploaded and processed successfully", "document_id": document.document_id, "chunk_eval": chunk_stats}, status=status.HTTP_201_CREATED)

        except Exception as e:
            # Mark document as failed in MongoDB
            try:
                self.mongo_repo.update_document_status(document.document_id if 'document' in locals() else file_hash, status="failed")
            except Exception:
                pass
            
            # Clean up the temporary file in case of an error
            if default_storage.exists(saved_path):
                default_storage.delete(saved_path)
            return Response({"error": str(e)}, status=500)

    def extract_text_and_metadata_from_pdf(self, file_path, document_id):
        """Extract text, metadata, and hierarchy from a PDF."""
        raw_blocks = []
        
        # Extract hierarchy information
        sections, document_title = extract_pdf_hierarchy(file_path)
        total_pages = 0
        
        with fitz.open(file_path) as pdf:
            total_pages = len(pdf)
            for page_number, page in enumerate(pdf, start=1):
                text = page.get_text()
                raw_blocks.append({
                    "text": text,
                    "page": page_number - 1,  # 0-indexed for consistency
                    "document_id": document_id
                })
        
        # Build page-to-hierarchy mapping
        structure_map = build_page_hierarchy_map(sections, total_pages) if sections else None
        
        # Attach hierarchy info to raw blocks
        for block in raw_blocks:
            page = block.get("page", 0)
            if structure_map and page in structure_map:
                hier = structure_map[page]
                block["section_title"] = hier.get("section_title")
                block["section_path"] = hier.get("section_path", [])
                block["hierarchy_level"] = hier.get("hierarchy_level", 0)
                block["parent_section"] = hier.get("parent_section")
            block["document_title"] = document_title
        
        # Store metadata for later use
        self._last_structure_map = structure_map
        self._last_document_title = document_title
        
        return raw_blocks

    def extract_text_and_metadata_from_docx(self, file_path, document_id):
        """Extract text, metadata, and hierarchy from a DOCX file."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        
        # Extract hierarchy information
        sections, document_title = extract_docx_hierarchy(file_path)
        
        raw_blocks = []
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if paragraph.text.strip():  # Skip empty paragraphs
                raw_blocks.append({
                    "text": paragraph.text,
                    "section": f"Paragraph {i}",
                    "page": i // 30,  # Approximate page number
                    "document_id": document_id
                })
        
        # Build structure map (using paragraph index as pseudo-page)
        total_pages = (len(doc.paragraphs) // 30) + 1
        structure_map = build_page_hierarchy_map(sections, total_pages) if sections else None
        
        # Attach hierarchy info to raw blocks
        for block in raw_blocks:
            page = block.get("page", 0)
            if structure_map and page in structure_map:
                hier = structure_map[page]
                block["section_title"] = hier.get("section_title")
                block["section_path"] = hier.get("section_path", [])
                block["hierarchy_level"] = hier.get("hierarchy_level", 0)
                block["parent_section"] = hier.get("parent_section")
            block["document_title"] = document_title
        
        # Store metadata for later use
        self._last_structure_map = structure_map
        self._last_document_title = document_title
        
        return raw_blocks

    def extract_text_and_metadata_from_txt(self, file_path, document_id):
        """Extract text and metadata from a TXT file."""
        raw_blocks = []
        with open(file_path, 'r', encoding='utf-8') as file:
            for i, line in enumerate(file, start=1):
                if line.strip():  # Skip empty lines
                    raw_blocks.append({
                        "text": line.strip(),
                        "section": f"Line {i}",
                        "page": i // 50,  # Approximate page (50 lines per page)
                        "document_id": document_id
                    })
        
        # TXT files typically don't have hierarchy
        self._last_structure_map = None
        self._last_document_title = None
        
        return raw_blocks


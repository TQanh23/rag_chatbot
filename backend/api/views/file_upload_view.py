from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from backend.utils.compute_file_hash import compute_file_hash
from documents.models import Document
from django.core.files.storage import default_storage
from backend.utils.qdrant_client import get_qdrant_client
from backend.utils.embeddings import HuggingfaceEmbeddingsModel
from qdrant_client.http.models import PointStruct, VectorParams, Distance, Filter, FieldCondition, MatchValue
from transformers import AutoTokenizer
import os
import uuid
import nltk
import fitz
from django.conf import settings
import logging
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from backend.utils.qdrant_client import QdrantClient

# Chunking and evaluation
from ..tools.chunk_eval import eval_chunks
from ..tools.export_chunks import write_chunks_csv

# MongoDB imports
from backend.utils.mongo_repository import MongoRepository
from backend.utils.mongo_models import MongoDocument, MongoChunk

logger = logging.getLogger(__name__)

@csrf_exempt
@require_http_methods(["POST"])
def upload_document(request):
    """
    Handle document upload and processing.
    """
    try:
        force = request.GET.get('force', 'false').lower() == 'true'
        
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        qdrant_client = QdrantClient()
        
        # If force=true, delete existing collection first
        if force:
            try:
                collection_name = os.environ.get('QDRANT_COLLECTION', 'test_collection')
                qdrant_client.delete_collection(collection_name)
                logger.info(f"Deleted existing collection: {collection_name}")
            except Exception as e:
                logger.warning(f"Could not delete collection (may not exist): {e}")
        
        # Now initialize/create collection
        qdrant_client.init_collection()
        
        # Process the file
        chunks = chunk_document(file)
        
        if not chunks:
            return JsonResponse({'error': 'No content extracted from file'}, status=400)
        
        # Generate embeddings and store in Qdrant
        texts = [chunk['text'] for chunk in chunks]
        embeddings = qdrant_client.embed_texts(texts)
        
        # Upload to vector DB
        qdrant_client.upload_vectors(embeddings, chunks)
        
        logger.info(f"Successfully uploaded document: {file.name} with {len(chunks)} chunks")
        
        return JsonResponse({
            'message': 'Document uploaded successfully',
            'file_name': file.name,
            'chunks_count': len(chunks)
        })
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)

class FileUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
    embeddings_model = HuggingfaceEmbeddingsModel('all-MiniLM-L6-v2')  # Initialize once
    mongo_repo = MongoRepository()  # MongoDB repository
    
    def delete_points_for_document(self, client, collection, document_id):
        """Delete all vector points for a specific document_id"""
        flt = Filter(must=[FieldCondition(key="document_id", match=MatchValue(value=document_id))])
        client.delete(collection_name=collection, points_selector=flt)
        logger.info(f"Deleted existing vector points for document_id: {document_id}")

    def post(self, request, *args, **kwargs):
        # Ensure NLTK 'punkt' is available
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')

        file = request.FILES.get('file')
        if not file:
            return Response(
                {"error": "No file uploaded. Make sure you use Body -> form-data and key name 'file' (type: File)."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Log file details
        logger.debug(f"Received file: name={file.name}, content_type={getattr(file, 'content_type', None)}, size={file.size}")

        # Basic validation
        allowed = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
        ]
        if file.content_type not in allowed:
            return Response(
                {"error": "Unsupported file type", "content_type": file.content_type},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Save the file temporarily
        saved_path = default_storage.save(f"uploads/{file.name}", file)
        logger.debug(f"Saved uploaded file to: {saved_path}")

        # Compute hash from the saved file
        saved_file_path = default_storage.path(saved_path) if hasattr(default_storage, 'path') else os.path.join(settings.MEDIA_ROOT, saved_path)
        with open(saved_file_path, "rb") as f:
            file_hash = compute_file_hash(f)

        # Check for force parameter - allows reingestion of the same file
        force = request.query_params.get("force") == "true"
        
        # Check for duplicates in Django DB
        existing_doc = Document.objects.filter(hash=file_hash).first()
        
        # Also check MongoDB
        existing_mongo_doc = self.mongo_repo.get_document_by_hash(file_hash) if not force else None
        
        if existing_doc and not force:
            default_storage.delete(saved_path)  # Clean up the temporary file
            return Response({
                'document_id': existing_doc.document_id,
                'filename': existing_doc.filename,
                'mimetype': existing_doc.mimetype,
                'size': existing_doc.size,
                'hash': existing_doc.hash,
                'message': 'Document already exists'
            })

        # If forcing, reuse the same document_id (hash) and delete existing points
        if existing_doc and force:
            client = get_qdrant_client()
            collection_name = settings.QDRANT_COLLECTION
            # Remove any leftover points (safe even if collection doesn't exist)
            try:
                self.delete_points_for_document(client, collection_name, existing_doc.document_id)
            except Exception as e:
                logger.warning(f"Could not delete existing points: {str(e)}")

            # Delete existing chunks in MongoDB
            self.mongo_repo.delete_chunks_by_document(existing_doc.document_id)
            
            # Reset counters on the existing document row
            existing_doc.token_count = 0
            existing_doc.vector_dim = None
            existing_doc.save()
            document = existing_doc
            logger.info(f"Force reingestion of document: {document.document_id}")
        else:
            # Save metadata (new document)
            document = Document.objects.create(
                document_id=file_hash,
                filename=file.name,
                mimetype=file.content_type,
                size=file.size,
                hash=file_hash,
            )
        
        # Create MongoDB document entry (will be updated after processing)
        mongo_doc = MongoDocument(
            document_id=file_hash,
            filename=file.name,
            content_hash=file_hash,
            mimetype=file.content_type,
            size_bytes=file.size,
            status="processing"
        )
        self.mongo_repo.save_document(mongo_doc)

        try:
            # Extract text and metadata based on file type
            if file.content_type == 'application/pdf':
                raw_blocks = self.extract_text_and_metadata_from_pdf(saved_file_path, document.document_id)
            elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                raw_blocks = self.extract_text_and_metadata_from_docx(saved_file_path, document.document_id)
            elif file.content_type == 'text/plain':
                raw_blocks = self.extract_text_and_metadata_from_txt(saved_file_path, document.document_id)
            else:
                raise ValueError("Unsupported file type")

            # Clean and chunk text
            from ..tools.chunking import clean_and_chunk_text
            chunks = clean_and_chunk_text(raw_blocks, self.tokenizer)
            logger.debug(f"Chunks generated: {chunks}")
            logger.debug(f"Number of chunks: {len(chunks)}")

            # Evaluate chunking quality before upload
            try:
                chunk_stats = eval_chunks(chunks, self.tokenizer, max_tokens=512)
                logger.info(f"Chunk evaluation: {chunk_stats}")
            except Exception as eval_exc:
                logger.exception("Chunk evaluation failed")
                chunk_stats = {"error": str(eval_exc)}

            # Generate embeddings using Hugging Face model
            chunk_texts = [chunk['text'] for chunk in chunks]
            embeddings = self.embeddings_model.embed_texts(chunk_texts)  # Batch embedding
            logger.debug(f"Embeddings generated: {embeddings}")
            logger.debug(f"Number of embeddings: {len(embeddings)}")

            # Ensure Qdrant collection exists
            client = get_qdrant_client()
            collection_name = settings.QDRANT_COLLECTION
            
            # If force=true, delete the collection first
            force = request.query_params.get("force") == "true"
            if force:
                try:
                    client.delete_collection(collection_name)
                    logger.info(f"Deleted existing collection: {collection_name}")
                except Exception as e:
                    logger.warning(f"Could not delete collection (may not exist): {e}")
            
            # Now create/ensure collection exists
            try:
                client.get_collection(collection_name)
            except Exception:
                # Try to create the collection; if it already exists (409) ignore the error.
                try:
                    client.create_collection(
                        collection_name=collection_name,
                        vectors_config={
                            "default": VectorParams(
                                size=len(embeddings[0]),
                                distance=Distance.COSINE
                            )
                        }
                    )
                except Exception as e:
                    msg = str(e).lower()
                    if "already exists" in msg or "409" in msg or "collection" in msg and "already exists" in msg:
                        logger.info(f"Collection {collection_name} already exists (ignored): {e}")
                    else:
                        logger.exception(f"Failed to create collection {collection_name}")
                        raise

            # Upload chunks and embeddings to Qdrant
            points = [
                PointStruct(
                    id=str(uuid.uuid4()),  # Generate a valid UUID for each point
                    vector={"default": embedding},  # Use named vector with "default"
                    payload=chunk
                )
                for embedding, chunk in zip(embeddings, chunks)
            ]
            logger.debug(f"Points being upserted: {points}")
            logger.debug(f"Number of points: {len(points)}")

            response = client.upsert(
                collection_name=collection_name,
                points=points
            )
            logger.debug(f"Upsert response: {response}")

            # Save chunks to MongoDB in smaller batches
            mongo_chunks = []
            for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                mongo_chunk = MongoChunk(
                    chunk_id=chunk.get('chunk_id', f"{document.document_id}::chunk_{idx}"),
                    document_id=document.document_id,
                    text=chunk['text'],
                    embedding=embedding,
                    page=chunk.get('page'),
                    end_page=chunk.get('end_page'),
                    section=chunk.get('section'),
                    order_index=chunk.get('order_index', idx),
                    text_len=len(chunk['text']),
                    metadata=chunk.get('metadata', {}),
                )
                mongo_chunks.append(mongo_chunk)
            
            # Bulk save chunks to MongoDB in batches of 100 to avoid BSON size limits
            batch_size = 100
            chunks_saved = 0
            for i in range(0, len(mongo_chunks), batch_size):
                batch = mongo_chunks[i:i + batch_size]
                try:
                    # Primary attempt: use repository helper
                    batch_count = self.mongo_repo.save_chunks_batch(batch)
                    # If repo returns falsy/0, fall back to direct insert into collection
                    if not batch_count:
                        logger.warning(f"repo.save_chunks_batch returned {batch_count} for batch {i // batch_size + 1}, falling back to direct insert")
                        docs = []
                        for c in batch:
                            docs.append({
                                "chunk_id": getattr(c, "chunk_id", None),
                                "document_id": getattr(c, "document_id", None),
                                "text": getattr(c, "text", None),
                                "embedding": getattr(c, "embedding", None),
                                "page": getattr(c, "page", None),
                                "end_page": getattr(c, "end_page", None),
                                "section": getattr(c, "section", None),
                                "order_index": getattr(c, "order_index", None),
                                "text_len": getattr(c, "text_len", None),
                                "metadata": getattr(c, "metadata", {}),
                            })
                        try:
                            result = self.mongo_repo.db["chunks"].insert_many(docs)
                            batch_count = len(result.inserted_ids)
                            logger.info(f"Fallback: inserted {batch_count} chunks directly into 'chunks' collection")
                        except Exception as ie:
                            logger.error(f"Fallback insert_many failed for batch {i // batch_size + 1}: {ie}")
                            batch_count = 0
                    chunks_saved += batch_count
                    logger.info(f"Saved batch {i // batch_size + 1}: {batch_count} / {len(batch)} chunks to MongoDB (total: {chunks_saved})")
                except Exception as e:
                    logger.exception(f"Failed to save batch {i // batch_size + 1}: {str(e)}")

            logger.info(f"✓ Total saved: {chunks_saved} chunks to MongoDB")

            # Update document status and store counts (Django DB)
            document.token_count = sum(len(self.tokenizer.encode(chunk['text'], add_special_tokens=False)) for chunk in chunks)
            document.vector_dim = len(embeddings[0])
            document.save()
            
            # Update MongoDB document with final status
            self.mongo_repo.update_document_status(
                document.document_id,
                status="processed",
                token_count=document.token_count,
                vector_dim=document.vector_dim,
                num_chunks=len(chunks)
            )

            # Write per-document CSV of chunks to MEDIA_ROOT/exports/{document_id}_chunks.csv
            try:
                csv_rel = f"exports/{document.document_id}_chunks.csv"
                csv_abs = os.path.join(settings.MEDIA_ROOT, csv_rel)
                # ensure exports directory exists
                os.makedirs(os.path.dirname(csv_abs), exist_ok=True)
                write_chunks_csv(chunks, csv_abs)
                logger.info(f"Wrote per-document chunks CSV: {csv_rel}")
            except Exception as csv_exc:
                logger.exception(f"Failed to write per-document CSV for {document.document_id}: {str(csv_exc)}")

            # Clean up the temporary file
            default_storage.delete(saved_path)

            # return chunk evaluation stats for quick validation (remove in production)
            return Response({"message": "File uploaded and processed successfully", "document_id": document.document_id, "chunk_eval": chunk_stats}, status=status.HTTP_201_CREATED)

        except Exception as e:
            # Mark document as failed in MongoDB
            try:
                self.mongo_repo.update_document_status(file_hash, status="failed")
            except Exception:
                pass
            
            # Clean up the temporary file in case of an error
            if default_storage.exists(saved_path):
                default_storage.delete(saved_path)
            return Response({"error": str(e)}, status=500)

    def extract_text_and_metadata_from_pdf(self, file_path, document_id):
        """Extract text and metadata from a PDF."""
        raw_blocks = []
        with fitz.open(file_path) as pdf:
            for page_number, page in enumerate(pdf, start=1):
                text = page.get_text()
                raw_blocks.append({
                    "text": text,
                    "page": page_number,
                    "document_id": document_id
                })
        return raw_blocks

    def extract_text_and_metadata_from_docx(self, file_path, document_id):
        """Extract text and metadata from a DOCX file."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        raw_blocks = []
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if paragraph.text.strip():  # Skip empty paragraphs
                raw_blocks.append({
                    "text": paragraph.text,
                    "section": f"Paragraph {i}",
                    "document_id": document_id
                })
        return raw_blocks

    def extract_text_and_metadata_from_txt(self, file_path, document_id):
        """Extract text and metadata from a TXT file."""
        raw_blocks = []
        with open(file_path, 'r', encoding='utf-8') as file:
            for i, line in enumerate(file, start=1):
                if line.strip():  # Skip empty lines
                    raw_blocks.append({
                        "text": line.strip(),
                        "section": f"Line {i}",
                        "document_id": document_id
                    })
        return raw_blocks


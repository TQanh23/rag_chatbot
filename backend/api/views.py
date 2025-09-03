from django.shortcuts import render
from httpcore import request
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
import fitz  # PyMuPDF
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, Distance, VectorParams
from django.conf import settings
import os
from backend.utils.qdrant_client import get_qdrant_client
from django.http import JsonResponse
from django.core.files.storage import default_storage
from documents.models import Document
from backend.utils.compute_file_hash import compute_file_hash  # Adjusted the import path to match the project structure
from backend.utils.embeddings import HuggingfaceEmbeddingsModel
import uuid
class AskView(APIView):
    def post(self, request):
        from langdetect import detect

        # Step 1: Normalize question
        question = request.data.get("question", "").strip()
        if not question:
            return Response({"error": "Question is required."}, status=status.HTTP_400_BAD_REQUEST)

        # Detect language (optional, based on your requirements)
        language = detect(question)
        print(f"Detected language: {language}")

        # Step 2: Embed the question
        embedding_model = HuggingfaceEmbeddingsModel('all-MiniLM-L6-v2')  # Replace with Gemini if needed
        question_embedding = embedding_model.embed_texts([question])[0]

        # Step 3: Retrieve from Qdrant
        client = get_qdrant_client()
        collection_name = settings.QDRANT_COLLECTION

        try:
            search_result = client.search(
                collection_name=collection_name,
                query_vector=("default", question_embedding),  # Use named vector format
                limit=20,
                with_payload=True,
                score_threshold=0.5
            )
            print(f"Query vector shape: {len(question_embedding)}")
        except Exception as e:
            return Response({"error": f"Error during Qdrant search: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Step 4: Format and return results
        results = [
            {
                "id": result.id,
                "score": result.score,
                "payload": result.payload
            }
            for result in search_result
        ]

        return Response({"results": results}, status=status.HTTP_200_OK)


class FileUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        # DEBUG: show what the server receives
        print(f"Request headers: {dict(request.headers)}")
        print(f"Request data keys: {list(request.data.keys())}")
        print(f"Request files keys: {list(request.FILES.keys())}")

        file = request.FILES.get('file')
        if not file:
            return Response(
                {"error": "No file uploaded. Make sure you use Body -> form-data and key name 'file' (type: File)."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Log file details
        print(f"Received file: name={file.name}, content_type={getattr(file, 'content_type', None)}, size={file.size}")

        # Basic validation
        allowed = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
        ]
        if file.content_type not in allowed:
            return Response(
                {"error": "Unsupported file type", "content_type": file.content_type},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Save temporarily (for debugging) — adjust to your real save/processing flow
        saved_path = default_storage.save(f"uploads/{file.name}", file)
        print(f"Saved uploaded file to: {saved_path}")

        # Compute hash
        file.seek(0)  # Reset file pointer to the beginning before hashing
        file_hash = compute_file_hash(file)

        # Check for duplicates
        existing_doc = Document.objects.filter(hash=file_hash).first()
        if existing_doc:
            return JsonResponse({
                'document_id': existing_doc.document_id,
                'filename': existing_doc.filename,
                'mimetype': existing_doc.mimetype,
                'size': existing_doc.size,
                'hash': existing_doc.hash,
            })

        # Save file
        file_path = default_storage.save(file.name, file)
        # Ensure the file path is absolute
        if not os.path.isabs(file_path):
            file_path = os.path.join(settings.MEDIA_ROOT, file_path)

        # Save metadata
        document = Document.objects.create(
            document_id=file_hash,
            filename=file.name,
            mimetype=file.content_type,
            size=file.size,
            hash=file_hash,
        )
        chunks = []  # Ensure chunks is always defined
        try:
            if file.content_type == 'application/pdf':
                raw_blocks = self.extract_text_and_metadata_from_pdf(file_path, document.document_id)
            elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                raw_blocks = self.extract_text_and_metadata_from_docx(file_path, document.document_id)
            elif file.content_type == 'text/plain':
                raw_blocks = self.extract_text_and_metadata_from_txt(file_path, document.document_id)
            else:
                raise ValueError("Unsupported file type")

            # Clean and chunk text
            chunks = self.clean_and_chunk_text(raw_blocks)

            # Generate embeddings using Hugging Face model
            chunk_texts = [chunk['text'] for chunk in chunks]
            model = HuggingfaceEmbeddingsModel('all-MiniLM-L6-v2')
            # Debug: Print chunk texts
            print(f"Chunk texts: {chunk_texts}")
            print(f"Number of chunks: {len(chunk_texts)}")

            # Generate embeddings
            embeddings = model.embed_texts(chunk_texts)  # Batch embedding

            # Debug: Print embedding details
            for i, embedding in enumerate(embeddings):
                print(f"Embedding {i}: Type: {type(embedding)}, Length: {len(embedding)}")

            # Ensure Qdrant collection exists
            client = get_qdrant_client()
            collection_name = settings.QDRANT_COLLECTION
            #     client.create_collection(
            #         collection_name=collection_name,
            #         vectors_config=VectorParams(size=len(embeddings[0]), distance="Cosine")
            #     )
            try:
                # Check if the collection exists
                collection_info = client.get_collection(collection_name)
                print(f"Collection '{collection_name}' exists: {collection_info}")
            except Exception:
                # Create the collection if it doesn't exist
                client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=len(embeddings[0]),  # Dimension of the embedding model
                        distance=Distance.COSINE  # Use cosine similarity
                    )
                )

            # points = [
            #     PointStruct(
            #         id=str(uuid.uuid4()),
            #         vector=embedding,  # No need to call .tolist() if embedding is already a list
            #         payload={
            #             "document_id": chunk["document_id"],
            #             "page": chunk.get("page"),
            #             "section": chunk.get("section"),
            #             "order_index": chunk["order_index"],
            #             "text": chunk["text"],
            #             "filename": document.filename,
            #             "hash": document.hash,
            #             "created_at": document.uploaded_at.isoformat()
            #         }
            #     )
            #     for chunk, embedding in zip(chunks, embeddings)
            # ]
            # for chunk, embedding in zip(chunks, embeddings):
            #     print(f"Embedding: {embedding}")
            #     print(f"Type: {type(embedding)}")
            #     print(f"Length: {len(embedding)}")
            # client.upsert(collection_name=collection_name, points=points)
            # Ensure proper point structure with vectors
                        # Ensure proper point structure with named vectors
            points = []
            for chunk, embedding in zip(chunks, embeddings):
                # Debug: Print embedding info
                print(f"Embedding type: {type(embedding)}, length: {len(embedding)}")
            
                points.append(
                    PointStruct(
                        id=str(uuid.uuid4()),  # Generate a unique ID
                        vector={"default": embedding},  # Use named vector with "default"
                        payload={
                            "document_id": chunk["document_id"],
                            "page": chunk.get("page"),
                            "section": chunk.get("section"),
                            "order_index": chunk["order_index"],
                            "text": chunk["text"],
                            "filename": document.filename,
                            "hash": document.hash,
                            "created_at": document.uploaded_at.isoformat()
                        }
                    )
                )
            
            # Upsert points to Qdrant
            response = client.upsert(
                collection_name=settings.QDRANT_COLLECTION,
                points=points
            )
            
            # Debug: Print upsert response
            print(f"Upsert response: {response}")

            # Update document status and store counts
            from transformers import AutoTokenizer
            tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
            document.token_count = sum(len(tokenizer.encode(chunk['text'], add_special_tokens=False)) for chunk in chunks)
            document.vector_dim = len(embeddings[0])
            document.save()
            document.vector_dim = len(embeddings[0])
            document.save()

            # Clean up the temporary file in a storage-agnostic way
            default_storage.delete(saved_path)

            return JsonResponse({"message": "File uploaded and processed successfully", "document_id": document.document_id}, status=201)

        except Exception as e:
            # Clean up the temporary file in case of an error
            if default_storage.exists(file_path):
                default_storage.delete(file_path)
            return JsonResponse({"error": str(e)}, status=500)

    def extract_text_from_pdf(self, file_path):
        """Extract text from a PDF using PyMuPDF."""
        text = ""
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text()
        return text

    def extract_text_from_docx(self, file_path):
        """Extract text from a DOCX file."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    def extract_text_from_txt(self, file_path):
        """Extract text from a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    def extract_text_and_metadata_from_pdf(self, file_path, document_id):
        """Extract text and metadata from a PDF."""
        raw_blocks = []
        with fitz.open(file_path) as pdf:
            for page_number, page in enumerate(pdf, start=1):
                text = page.get_text()
                raw_blocks.append({
                    "text": text,
                    "page": page_number,
                    "document_id": document_id
                })
        return raw_blocks

    def extract_text_and_metadata_from_docx(self, file_path, document_id):
        """Extract text and metadata from a DOCX file."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        raw_blocks = []
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if paragraph.text.strip():  # Skip empty paragraphs
                raw_blocks.append({
                    "text": paragraph.text,
                    "section": f"Paragraph {i}",
                    "document_id": document_id
                })
        return raw_blocks

    def extract_text_and_metadata_from_txt(self, file_path, document_id):
        """Extract text and metadata from a TXT file."""
        raw_blocks = []
        with open(file_path, 'r', encoding='utf-8') as file:
            for i, line in enumerate(file, start=1):
                if line.strip():  # Skip empty lines
                    raw_blocks.append({
                        "text": line.strip(),
                        "section": f"Line {i}",
                        "document_id": document_id
                    })
        return raw_blocks
    def chunk_text(self, text, chunk_size=500):
        """Split text into smaller chunks."""
        words = text.split()
        chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]
        return chunks

    def clean_and_chunk_text(self, raw_blocks, chunk_size=1200, overlap=200):
        """Clean and chunk text into overlapping segments."""
        chunks = []
        chunk_id = 0

        for block in raw_blocks:
            # Normalize whitespace
            text = ' '.join(block['text'].split())

            # Remove headers/footers (example logic, can be customized)
            lines = text.split('\n')
            if len(lines) > 2:  # Assume header/footer if more than 2 lines
                text = '\n'.join(lines[1:-1])

            # Recursive chunking
            start = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                chunk_text = text[start:end]

                # Add chunk metadata
                chunks.append({
                    "chunk_id": f"chunk_{chunk_id}",
                    "text": chunk_text,
                    "document_id": block["document_id"],
                    "page": block.get("page"),
                    "section": block.get("section"),
                    "order_index": chunk_id
                })

                chunk_id += 1
                start += chunk_size - overlap

        return chunks

class InitQdrantView(APIView):
    def get(self, request):
        client = get_qdrant_client()

        try:
            client.get_collection(settings.QDRANT_COLLECTION)
            return Response({"message": f"Qdrant collection '{settings.QDRANT_COLLECTION}' already exists."})
        except Exception:
            client.recreate_collection(
                collection_name=settings.QDRANT_COLLECTION,
                vectors_config={
                    "default": VectorParams(
                        size=384,  # Adjust based on your embedding model
                        distance=Distance.COSINE  # Use "COSINE", "EUCLID", or "DOT" based on your use case
                    )
                }
            )
            return Response({"message": f"Qdrant collection '{settings.QDRANT_COLLECTION}' created successfully."})
class DeleteQdrantCollectionView(APIView):
    def post(self, request):
        client = get_qdrant_client()
        try:
            client.delete_collection(settings.QDRANT_COLLECTION)
            return Response({"message": "Collection deleted."})
        except Exception as e:
            return Response({"error": str(e)}, status=500)